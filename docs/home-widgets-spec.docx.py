#!/usr/bin/env python3
"""Generate Home Widgets Specification .docx"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ===== TITLE =====
title = doc.add_heading('Спецификация виджетов домашней страницы Evo App', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    f'Дата составления: {datetime.date.today().strftime("%d.%m.%Y")}\n'
    'Проект: Evo App (evo-app)\n'
    'Страница: Home (/pages/Home.tsx)\n'
    'Фреймворк: React + TypeScript, Hono RPC, TanStack Query, Tailwind CSS, Framer Motion'
).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ===== HELPERS =====
def add_widget_heading(name, subtitle=""):
    h = doc.add_heading(name, level=1)
    if subtitle:
        doc.add_paragraph(subtitle).italic = True
    return h

def add_section(heading_text):
    return doc.add_heading(heading_text, level=2)

def add_prop_table(props):
    """props: list of (name, type, required, description)"""
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = 'Пропс'
    hdr[1].text = 'Тип'
    hdr[2].text = 'Обяз.'
    hdr[3].text = 'Описание'
    for row_data in props:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
    doc.add_paragraph()

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_state_table(states):
    """states: list of (state_name, component, description)"""
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Состояние'
    hdr[1].text = 'Компонент'
    hdr[2].text = 'Описание'
    for row_data in states:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
    doc.add_paragraph()

# ═══════════════════════════════════════════
# 1. DAILY BRIEFING
# ═══════════════════════════════════════════
add_widget_heading('1. DailyBriefing — Ежедневный брифинг', 'Виджет приветствия с персональной аналитикой')

doc.add_paragraph(
    'Расположен в самом верху домашней страницы. Показывает персональное приветствие '
    'по времени суток, имя сотрудника, текущий магазин и ключевые метрики эффективности.'
)

add_section('Источники данных')
doc.add_paragraph('• useEmployeeNameAndUuid() — имя и UUID сотрудника из /api/employees/by-last-name-uuid')
doc.add_paragraph('• useGetReportAndPlan(true) — данные отчёта и плана (fetchFinancialForToday + fetchPlanForToday)')
doc.add_paragraph('• useWorkingByShops() — информация о работающих сегодня сотрудниках по магазинам')
doc.add_paragraph('• useSellerEffectiveness({ period: 30 }) — метрики эффективности продавца за 30 дней')
doc.add_paragraph('• isTelegramMiniApp() — проверка запуска внутри Telegram Mini App')

add_section('Состояния')
add_state_table([
    ('Загрузка', 'BriefingSkeleton', 'Анимированный плейсхолдер: градиентный блок сине-индиго, пульсирующие полосы'),
    ('Данные (сотрудник найден)', 'Полный брифинг', 'Приветствие + имя + магазин + план + метрики эффективности'),
    ('Нет сотрудника', 'Упрощённый брифинг', 'Только приветствие без персональных данных'),
    ('Telegram Mini App', 'Компактный вид', 'Уменьшенные отступы, текст мельче (text-sm)'),
])

add_section('Стили')
doc.add_paragraph('• Фон: градиент from-blue-500 to-indigo-600 (светлая тема), from-blue-600 to-indigo-700 (тёмная)')
doc.add_paragraph('• Скругление: rounded-xl (12px)')
doc.add_paragraph('• Тень: shadow-lg')
doc.add_paragraph('• Отступы: p-4 (16px), mb-4 (16px снизу)')
doc.add_paragraph('• Текст: белый (text-white), разный размер (text-xs / text-sm / text-2xl)')
doc.add_paragraph('• Анимация: Framer Motion — появление снизу (initial={{ opacity: 0, y: -12 }}, animate={{ opacity: 1, y: 0 }})')
doc.add_paragraph('• Иконки: MapPin (магазин), Target (план), TrendingUp/TrendingDown (тренд) из lucide-react')
doc.add_paragraph('• Бейджи метрик: скруглённые (rounded-lg), полупрозрачный фон (bg-white/10), отступ px-2.5 py-1')

add_section('Приветствие по времени')
add_code(
    '06:00–11:59 → "Доброе утро"\n'
    '12:00–17:59 → "Добрый день"\n'
    '18:00–23:59 → "Добрый вечер"\n'
    '00:00–05:59 → "Доброй ночи"'
)

add_section('Метрики эффективности (при наличии сотрудника)')
doc.add_paragraph('• Текущий магазин: иконка MapPin, название магазина')
doc.add_paragraph('• План на сегодня: иконка Target, сумма в ₽ (формат: ≥1M → "1.2M", ≥1k → "12k", иначе число)')
doc.add_paragraph('• Рейтинг продавца: значок медали (🥇🥈🥉) + "#N из M"')
doc.add_paragraph('• Тренд выручки: TrendingUp/TrendingDown, ₽/день, цвет: зелёный/красный')

doc.add_paragraph('Формат чисел: ≥ 1 000 000 → "1.2M", ≥ 1 000 → "12k", иначе целое число', style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════
# 2. DATE FILTER
# ═══════════════════════════════════════════
add_widget_heading('2. DateFilter — Фильтр дат', 'Переключение периода отображения данных')

doc.add_paragraph(
    'Расположен под DailyBriefing. Три кнопки: Сегодня / Вчера / Период (с календарём). '
    'При переключении все виджеты на странице обновляют данные за выбранный период.'
)

add_section('Пропсы')
add_prop_table([
    ('value', 'DateFilterValue { since, until, dateMode }', 'Да', 'Текущий выбранный период'),
    ('onChange', '(v: DateFilterValue) => void', 'Да', 'Колбэк при смене периода'),
])

add_section('Режимы')
doc.add_paragraph('• today — сегодняшняя дата (since = until = YYYY-MM-DD)')
doc.add_paragraph('• yesterday — вчерашняя дата')
doc.add_paragraph('• period — произвольный диапазон через Popover с Calendar (react-day-picker)')

add_section('Стили')
doc.add_paragraph('• Сетка: grid grid-cols-3 gap-2, mb-4')
doc.add_paragraph('• Кнопка активна: border-blue-600 bg-blue-600 text-white')
doc.add_paragraph('• Кнопка неактивна: border-gray-300 bg-white text-gray-800 (светлая), dark:border-gray-700 dark:bg-gray-800 dark:text-gray-200')
doc.add_paragraph('• Скругление: rounded-lg, отступы: px-3 py-2')
doc.add_paragraph('• Календарь: Popover из @/components/ui, Calendar из react-day-picker')
doc.add_paragraph('• При выборе периода: кнопка показывает "MM-DD – MM-DD"')

add_section('Поведение')
doc.add_paragraph('• Календарь открывается по клику на «Период»')
doc.add_paragraph('• При клике «Сегодня» или «Вчера» — календарь закрывается, данные обновляются мгновенно')
doc.add_paragraph('• В режиме «period» — кнопки «Применить»/«Отмена» внутри Popover')

doc.add_page_break()

# ═══════════════════════════════════════════
# 3. REVENUE WIDGET
# ═══════════════════════════════════════════
add_widget_heading('3. RevenueWidget — Выручка', 'Карточка чистой выручки с детализацией по магазинам')

add_section('Пропсы')
add_prop_table([
    ('since', 'string (YYYY-MM-DD)', 'Да', 'Начало периода'),
    ('until', 'string (YYYY-MM-DD)', 'Да', 'Конец периода'),
    ('expanded', 'boolean', 'Да', 'Развёрнута ли детализация'),
    ('onToggle', '() => void', 'Да', 'Переключение разворота'),
])

add_section('Источники данных')
doc.add_paragraph('• useSalesData({ since, until, shopUuid }) — сырые данные продаж')
doc.add_paragraph('• useFilteredSalesData() — фильтрация по магазину/роли')
doc.add_paragraph('• useSalesCalculations() — вычисление netSales = totalSell - totalRefund')
doc.add_paragraph('• Для SUPERADMIN: shopUuid = undefined (все магазины)')
doc.add_paragraph('• Для обычного пользователя: shopUuid = текущий магазин из useCurrentWorkShop()')

add_section('Состояния')
add_state_table([
    ('Загрузка', 'SkeletonCard tone="blue"', 'Плейсхолдер с левой синей полосой (border-l-4 border-l-blue-500)'),
    ('Ошибка', 'div text-red-500', 'Текст ошибки'),
    ('Данные (свёрнуто)', 'RevenueCard', 'Градиентная карточка с суммой выручки'),
    ('Данные (развёрнуто)', 'RevenueDetailsAdmin / RevenueDetailsUser', 'Детализация по магазинам'),
])

add_section('RevenueCard — стили')
doc.add_paragraph('• Фон: градиент from-blue-500 to-blue-600 (светлая), dark:from-blue-600 dark:to-blue-700')
doc.add_paragraph('• Текст: белый (text-white), тень shadow-lg')
doc.add_paragraph('• Заголовок: «Выручка» (text-xs font-medium opacity-90) + иконка DollarSign')
doc.add_paragraph('• Сумма: text-2xl font-bold, формат через formatCurrency()')
doc.add_paragraph('• Прогресс-бар: абсолютно позиционирован снизу, h-1, bg-white/30, заполнение bg-white/60 на 70%')
doc.add_paragraph('• Ховер: scale: 1.03, y: -2 (Framer Motion whileHover)')
doc.add_paragraph('• Тап: scale: 0.98 (whileTap)')
doc.add_paragraph('• Развёрнуто: ring-2 ring-blue-500 scale-[1.01]')

add_section('RevenueDetailsAdmin — содержимое')
doc.add_paragraph('• Таблица по магазинам: Название | Продажи | Возвраты | Чистая | % возвратов | Чеков | Средний чек')
doc.add_paragraph('• Итоговая строка с суммарными показателями')
doc.add_paragraph('• Цветовая индикация: % возвратов > 10% — красный')
doc.add_paragraph('• Фон: bg-white dark:bg-gray-800, скругление rounded-xl, тень shadow')

doc.add_page_break()

# ═══════════════════════════════════════════
# 4. SALES TEMPO WIDGET
# ═══════════════════════════════════════════
add_widget_heading('4. SalesTempoWidget — Темп продаж', 'Сравнение с предыдущим периодом и почасовой график')

add_section('Пропсы')
add_prop_table([
    ('since', 'string', 'Да', 'Начало периода'),
    ('until', 'string', 'Да', 'Конец периода'),
    ('expanded', 'boolean', 'Да', 'Развёрнута ли детализация'),
    ('onToggle', '() => void', 'Да', 'Переключение'),
])

add_section('Логика сравнения')
doc.add_paragraph('• Вычисляется предыдущий период той же длины: prevSince/prevUntil')
doc.add_paragraph('• Delta% = ((current - previous) / previous) × 100')
doc.add_paragraph('• Дополнительно загружаются данные по аксессуарам (useAccessoriesSales) при развороте')

add_section('RevenueTempoCard — стили')
doc.add_paragraph('• Заголовок: «Темп продаж» + иконка Clock3')
doc.add_paragraph('• Дельта: зелёная при росте, красная при падении')
doc.add_paragraph('• Развёрнуто: ring-2 ring-slate-500 scale-[1.01]')
doc.add_paragraph('• Скелетон: SkeletonCard tone="indigo"')

add_section('RevenueTempoDetails — содержимое (500 строк кода)')
doc.add_paragraph('• Сводка: Выручка | Чеки | Средний чек | Возвраты | Изменение vs прошлый период')
doc.add_paragraph('• Таблица по магазинам с теми же метриками')
doc.add_paragraph('• Почасовой план-факт график: fetchRevenueHourlyPlanFact() → данные по часам')
doc.add_paragraph('• Два режима отображения: cumulative (накопленный) / hourly (по часам)')
doc.add_paragraph('• График: Recharts (LineChart, XAxis, YAxis, Tooltip, Legend)')
doc.add_paragraph('• Секция аксессуаров: доля аксессуаров в выручке, сравнение с целевым показателем')
doc.add_paragraph('• Целевой показатель доли аксессуаров настраивается (useTempoSettings)')

doc.add_page_break()

# ═══════════════════════════════════════════
# 5. FINANCE WIDGET
# ═══════════════════════════════════════════
add_widget_heading('5. FinanceWidget — Финансовый отчёт', 'Расходы, баланс наличных, кассовые операции')

add_section('Пропсы')
add_prop_table([
    ('since', 'string', 'Да', 'Начало периода'),
    ('until', 'string', 'Да', 'Конец периода'),
    ('expanded', 'boolean', 'Да', 'Развёрнут'),
    ('onToggle', '() => void', 'Да', 'Переключение'),
])

add_section('ExpensesCard — стили')
doc.add_paragraph('• Фон: градиент from-orange-500 to-orange-600, dark:from-orange-600 dark:to-orange-700')
doc.add_paragraph('• Заголовок: «Фин. отчёт» + иконка ShoppingCart')
doc.add_paragraph('• Превью (свёрнуто): топ-3 магазина по кассовому балансу с суммами')
doc.add_paragraph('• Если кассовый баланс = 0 у всех — показываются магазины с продажами')
doc.add_paragraph('• Счётчик скрытых магазинов: «+N» если больше трёх')
doc.add_paragraph('• Развёрнуто: ring-2 ring-orange-500 scale-[1.01]')
doc.add_paragraph('• Скелетон: SkeletonCard tone="orange"')

add_section('FinancialReportDetails — содержимое')
doc.add_paragraph('• Таблица: Магазин | Продажи | Возвраты | Расходы (CASH_OUTCOME) | Баланс')
doc.add_paragraph('• Итоговая строка')
doc.add_paragraph('• Каждое поле — с форматированием через formatCurrency()')
doc.add_paragraph('• Баланс = Продажи - Возвраты - Расходы')
doc.add_paragraph('• Отрицательный баланс подсвечивается красным')

doc.add_page_break()

# ═══════════════════════════════════════════
# 6. BEST SHOP WIDGET
# ═══════════════════════════════════════════
add_widget_heading('6. BestShopWidget — Лучший магазин', 'Лидер дня и недели по выручке')

add_section('Пропсы')
add_prop_table([
    ('since', 'string', 'Да', 'Начало периода'),
    ('until', 'string', 'Да', 'Конец периода'),
    ('dateMode', '"today" | "yesterday" | "period"', 'Да', 'Режим даты'),
    ('expanded', 'boolean', 'Да', 'Развёрнут'),
    ('onToggle', '() => void', 'Да', 'Переключение'),
])

add_section('Источники данных')
doc.add_paragraph('• useDashboardHomeInsights() — возвращает bestShop.dayLeader, bestShop.weekLeader, bestShop.dayRows, bestShop.weekRows')

add_section('BestShopCard — стили')
doc.add_paragraph('• Фон: градиент from-purple-500 to-purple-600, dark:from-purple-600 dark:to-purple-700')
doc.add_paragraph('• Заголовок: «Топ магазин» + иконка Award (Trophy)')
doc.add_paragraph('• Два режима (mode): "day" (лидер дня) / "week" (лидер недели)')
doc.add_paragraph('• Переключение режимов внутри карточки')
doc.add_paragraph('• Показывает: название магазина, чистую выручку, причину лидерства')
doc.add_paragraph('• Причина: «чек» (высокий средний чек), «трафик» (много покупателей), «конверсия»')
doc.add_paragraph('• Разница со вторым местом (gapToSecond)')
doc.add_paragraph('• Скелетон: SkeletonCard tone="purple"')

add_section('BestShopDetails — содержимое')
doc.add_paragraph('• Переключатель day/week')
doc.add_paragraph('• Рейтинговая таблица всех магазинов: # | Магазин | Выручка | Отрыв')
doc.add_paragraph('• Лидер выделен: золотой фон, иконка Award')
doc.add_paragraph('• Медали: 🥇 золото, 🥈 серебро, 🥉 бронза для топ-3')

doc.add_page_break()

# ═══════════════════════════════════════════
# 7. TOP PRODUCT WIDGET
# ═══════════════════════════════════════════
add_widget_heading('7. TopProductWidget — Топ продуктов', 'Самый продаваемый товар с детализацией')

add_section('Пропсы')
add_prop_table([
    ('since', 'string', 'Да', 'Начало периода'),
    ('until', 'string', 'Да', 'Конец периода'),
    ('expanded', 'boolean', 'Да', 'Развёрнут'),
    ('onToggle', '() => void', 'Да', 'Переключение'),
])

add_section('Режимы метрики (metricMode)')
doc.add_paragraph('• "revenue" — сортировка по выручке (по умолчанию)')
doc.add_paragraph('• "quantity" — сортировка по количеству')
doc.add_paragraph('• "margin" — сортировка по маржинальности (%)')

add_section('Фильтры возвратов (refundFilter)')
doc.add_paragraph('• "all" — все товары')
doc.add_paragraph('• "noRefunds" — только без возвратов')
doc.add_paragraph('• "highRefund" — товары с высокой долей возвратов (≥8%)')

add_section('TopProductCard — стили')
doc.add_paragraph('• Фон: градиент from-pink-500 to-pink-600, dark:from-pink-600 dark:to-pink-700')
doc.add_paragraph('• Заголовок: «Топ продукт» + иконка Package')
doc.add_paragraph('• Показывает: название продукта, выручку/количество/маржа (в зависимости от режима)')
doc.add_paragraph('• Тренд vs предыдущий период: TrendingUp (зелёный) / TrendingDown (красный)')
doc.add_paragraph('• Развёрнуто: ring-2 ring-pink-500')
doc.add_paragraph('• Скелетон: SkeletonCard tone="pink"')
doc.add_paragraph('• Пусто: EmptyTile с иконкой Package')

add_section('TopProductsDetails — содержимое')
doc.add_paragraph('• Таблица: # | Товар | Выручка | Количество | Маржа% | Возвраты%')
doc.add_paragraph('• Сортировка по выбранной метрике')
doc.add_paragraph('• Прогресс-бары для визуализации доли каждого товара')

doc.add_page_break()

# ═══════════════════════════════════════════
# 8. SELLER PERFORMANCE
# ═══════════════════════════════════════════
add_widget_heading('8. SellerPerformanceWidget — Топ продавцов', 'Топ-3 продавца сегодня (только SUPERADMIN)')

doc.add_paragraph('Виден только для SUPERADMIN. Компактный виджет с навигацией на полный отчёт.')

add_section('Источники данных')
doc.add_paragraph('• useSellerEffectiveness({ period: 1 }) — эффективность продавцов за сегодня')
doc.add_paragraph('• Фильтруются только активные (daysWorked ≥ 1), показываются топ-3')

add_section('Состояния')
add_state_table([
    ('Загрузка', 'Skeleton', '3 строки с пульсирующими аватарами и полосами текста'),
    ('Нет данных', 'Empty state', 'Иконка Trophy (opacity-30), текст «Нет данных за сегодня», кликабельно — переход на /evotor/seller-performance'),
    ('Данные', 'Топ-3', 'Карточка с тремя продавцами'),
])

add_section('Стили')
doc.add_paragraph('• Контейнер: bg-white dark:bg-gray-800, rounded-xl, shadow-sm, border, overflow-hidden')
doc.add_paragraph('• Заголовок: «Продавцы дня» + иконка ChevronRight → переход на полный отчёт')
doc.add_paragraph('• Места: 🥇 золото (amber-500), 🥈 серебро (slate-400), 🥉 бронза (amber-700)')
doc.add_paragraph('• Строка продавца: аватар (цветной круг с инициалами), имя (truncate), выручка (font-medium)')
doc.add_paragraph('• Аватары: цвет из палитры по индексу (amber, blue, emerald, violet, rose)')
doc.add_paragraph('• Ховер: bg-gray-50 dark:bg-gray-750')
doc.add_paragraph('• Тап: переход на /evotor/seller-performance')

doc.add_page_break()

# ═══════════════════════════════════════════
# 9. TODAY ALERTS
# ═══════════════════════════════════════════
add_widget_heading('9. TodayAlertsWidget — Алерты', 'Автоматические предупреждения на основе данных')

add_section('Источники данных')
doc.add_paragraph('• fetchFinancialForToday() — данные финансового отчёта за сегодня')
doc.add_paragraph('• buildTodayAlerts() — бизнес-логика генерации алертов')
doc.add_paragraph('• Модель: TodayAlertModel { type, title, message, iconKey }')

add_section('Типы алертов')
add_state_table([
    ('warning', 'Clock / AlertTriangle', 'Предупреждение (жёлтый/amber)'),
    ('danger', 'AlertTriangle', 'Критическое (красный/red)'),
    ('info', 'TrendingDown', 'Информационное (синий/blue)'),
])

doc.add_paragraph('Примеры алертов:')
doc.add_paragraph('• Падение выручки vs вчера > порога → "danger"')
doc.add_paragraph('• Магазин не открыт к 07:50 → "warning" + иконка Clock')
doc.add_paragraph('• Низкая доля вейпов → "info"')

add_section('Стили')
doc.add_paragraph('• Контейнер алерта: border-l-4, скругление rounded-lg, отступ p-3')
doc.add_paragraph('• warning: border-l-amber-500, bg-amber-500/5')
doc.add_paragraph('• danger: border-l-red-500, bg-red-500/5')
doc.add_paragraph('• info: border-l-blue-500, bg-blue-500/5')
doc.add_paragraph('• Иконка: w-4 h-4, цвет соответствует типу')
doc.add_paragraph('• Заголовок: text-sm font-medium')
doc.add_paragraph('• Сообщение: text-xs text-gray-500 dark:text-gray-400')
doc.add_paragraph('• Автообновление: refetchInterval 120_000 (2 минуты), staleTime 30_000')
doc.add_paragraph('• При отсутствии алертов: «✅ Всё в порядке» на зелёном фоне (bg-emerald-500/10)')

doc.add_page_break()

# ═══════════════════════════════════════════
# 10. STOCK HEALTH
# ═══════════════════════════════════════════
add_widget_heading('10. StockHealthWidget — Состояние склада', 
    'Мёртвый сток, низкий сток, закончившиеся товары (SUPERADMIN/ADMIN)')

doc.add_paragraph('Самый сложный виджет (500 строк). Показывает три категории проблемных товаров с возможностью перемещения между магазинами.')

add_section('Источники данных')
doc.add_paragraph('• useStockHealth(days) — загрузка данных о стоке')
doc.add_paragraph('• useStockTransfer() — мутация для перемещения товаров между магазинами')
doc.add_paragraph('• Типы: StockItem, OutOfStockItem, TransferRec, StockHealthData')

add_section('Три секции (разворачиваемые)')
doc.add_paragraph('1. Мёртвый сток (DeadStock) — товары без продаж за N дней (по умолчанию 14)')
doc.add_paragraph('2. Низкий сток (LowStock) — товары с остатком ниже порога')
doc.add_paragraph('3. Out of Stock — закончившиеся товары (были в продаже, сейчас 0)')

add_section('Фильтры')
doc.add_paragraph('• Выбор магазина: "all" или конкретный (ShopSelector)')
doc.add_paragraph('• Настройка дней для мёртвого стока: 14 (по умолчанию), можно менять (7, 14, 21, 30)')
doc.add_paragraph('• Фильтр «Только с перемещениями» — показывает товары, которые есть в других магазинах')

add_section('Стили')
doc.add_paragraph('• Контейнер: bg-white dark:bg-gray-800, rounded-xl, shadow, p-4')
doc.add_paragraph('• Заголовок: «Состояние склада» + кнопка экспорта в Excel')
doc.add_paragraph('• KPI-блоки (3 в ряд):')
doc.add_paragraph('  - Мёртвый сток: иконка Skull, количество + сумма (text-red-600)')
doc.add_paragraph('  - Низкий сток: иконка AlertTriangle, количество (text-amber-600)')
doc.add_paragraph('  - Out of Stock: иконка PackageX, количество (text-orange-600)')
doc.add_paragraph('• Каждая секция — разворачиваемая (ChevronDown/ChevronRight)')
doc.add_paragraph('• Таблица товаров: Название | Остаток | Последняя продажа | Сумма | Переместить')
doc.add_paragraph('• Кнопка «Переместить»: открывает диалог выбора магазина-получателя')
doc.add_paragraph('• Индикатор перемещения: зелёный бейдж «Можно переместить в X»')
doc.add_paragraph('• Пустое состояние: «✅ Мёртвого стока нет» / «Все товары продаются»')

add_section('Экспорт')
doc.add_paragraph('• Кнопка «Экспорт в Excel» в заголовке')
doc.add_paragraph('• Формирует URL для скачивания CSV/XLSX через backend')

doc.add_page_break()

# ═══════════════════════════════════════════
# 11. QUICK ACTIONS
# ═══════════════════════════════════════════
add_widget_heading('11. QuickActionsWidget — Быстрые действия', 'Сетка кнопок быстрого доступа к функциям')

add_section('Пропсы')
add_prop_table([
    ('employeeRole', 'string', 'Да', 'Роль сотрудника (CASHIER/ADMIN/SUPERADMIN)'),
])

add_section('Список действий (зависит от роли)')
doc.add_paragraph('Доступные действия определяются через getAvailableQuickActions(role):')
doc.add_paragraph('• Открытие смены (door_open) → /evotor/open-store — все роли')
doc.add_paragraph('• Мёртвый сток (package) → /evotor/dead-stock — все роли')
doc.add_paragraph('• Отчёт по продажам (file_text) → /evotor/sales-report — все роли')
doc.add_paragraph('• Эффективность продавцов (trending_up) → /evotor/seller-performance — SUPERADMIN')
doc.add_paragraph('• Админ-открытия (store) → /evotor/store-openings-admin — SUPERADMIN')
doc.add_paragraph('• Зарплатный отчёт (calculator) → /evotor/salary-report — SUPERADMIN')
doc.add_paragraph('• AI-ассистент (sparkles) → открывает Telegram бота — SUPERADMIN в Mini App')

add_section('Бейджи')
doc.add_paragraph('• На кнопках «Мёртвый сток» и «Низкий сток» — бейджи с количеством')
doc.add_paragraph('• Данные загружаются через useStockHealth(14)')
doc.add_paragraph('• Бейдж: красный круг (bg-red-500), белый текст, позиционирован справа-сверху')

add_section('Стили')
doc.add_paragraph('• Заголовок: «Быстрые действия» (text-sm font-semibold)')
doc.add_paragraph('• Сетка: grid grid-cols-2 lg:grid-cols-4 gap-3')
doc.add_paragraph('• Кнопка: bg-white dark:bg-gray-800, rounded-xl, p-4, shadow-sm, border')
doc.add_paragraph('• Иконка: w-6 h-6, цветная (amber/blue/emerald/purple/rose)')
doc.add_paragraph('• Подпись: text-xs text-gray-600 dark:text-gray-400, mt-1')
doc.add_paragraph('• Ховер: bg-gray-50 dark:bg-gray-750, shadow-md')
doc.add_paragraph('• Тап: active:scale-[0.97]')

doc.add_page_break()

# ═══════════════════════════════════════════
# 12. WIDGET ERROR BOUNDARY
# ═══════════════════════════════════════════
add_widget_heading('12. WidgetErrorBoundary — Обработка ошибок', 'Изоляция сбоев отдельных виджетов')

doc.add_paragraph('Каждый виджет на домашней странице обёрнут в WidgetErrorBoundary. При ошибке в виджете он не роняет всю страницу, а показывает заглушку с кнопкой повтора.')

add_section('Пропсы')
add_prop_table([
    ('children', 'ReactNode', 'Да', 'Содержимое виджета'),
    ('fallback', 'ReactNode', 'Нет', 'Кастомная заглушка (опционально)'),
    ('onRetry', '() => void', 'Нет', 'Действие при повторе'),
])

add_section('Стили заглушки')
doc.add_paragraph('• Фон: bg-red-50 dark:bg-red-950/20')
doc.add_paragraph('• Граница: border border-red-200 dark:border-red-800')
doc.add_paragraph('• Скругление: rounded-xl')
doc.add_paragraph('• Заголовок: «⚠️ Ошибка загрузки виджета» (text-red-600 dark:text-red-400, text-sm)')
doc.add_paragraph('• Сообщение: текст ошибки (text-xs, max-h-20, overflow-hidden)')
doc.add_paragraph('• Кнопка: «🔄 Повторить» — bg-red-100 dark:bg-red-900/40, hover:bg-red-200')

add_section('Использование в Home.tsx')
add_code(
    '<ErrorBoundary variant="widget" name="Выручка">\n'
    '  <RevenueWidget ... />\n'
    '</ErrorBoundary>'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 13. ОБЩИЕ КОМПОНЕНТЫ
# ═══════════════════════════════════════════
add_widget_heading('13. Общие UI-компоненты', 'Переиспользуемые элементы')

add_section('SkeletonCard')
doc.add_paragraph('Плейсхолдер для карточек во время загрузки.')
doc.add_paragraph('• Пропсы: tone ("blue"|"orange"|"purple"|"pink"|"cyan"|"indigo")')
doc.add_paragraph('• Левая цветная полоса: border-l-4 border-l-{tone}-500')
doc.add_paragraph('• Внутри: пульсирующий квадрат (иконка) + две полосы текста')
doc.add_paragraph('• Анимация: animate-pulse')
doc.add_paragraph('• Фон: bg-white dark:bg-gray-800, rounded-xl, shadow')

add_section('EmptyTile')
doc.add_paragraph('Заглушка при отсутствии данных.')
doc.add_paragraph('• Иконка: серая (text-gray-300), opacity-60')
doc.add_paragraph('• Текст: «{title}» (text-xs text-gray-400)')
doc.add_paragraph('• Минимальная высота: min-h-[120px]')

add_section('LoadingTile')
doc.add_paragraph('Альтернативный плейсхолдер с иконкой.')
doc.add_paragraph('• Цвет иконки зависит от tone (через toneMap)')
doc.add_paragraph('• Анимация: animate-pulse')
doc.add_paragraph('• Текст: название загружаемого виджета')

add_section('AnimatedNumber')
doc.add_paragraph('Анимированное число (из /widgets/dashboard/ui/AnimatedNumber.tsx).')
doc.add_paragraph('• Плавное изменение значения через Framer Motion useSpring')
doc.add_paragraph('• Используется для анимации сумм выручки при переключении дат')

add_section('formatCurrency (утилита)')
doc.add_paragraph('Форматирование валюты:')
add_code(
    'formatCurrency(1234567) → "1 234 567"\n'
    'formatCurrency(0) → "0"'
)
doc.add_paragraph('Разделитель тысяч — пробел, без десятичных знаков.')

# ═══════════════════════════════════════════
# 14. СХЕМА ДАННЫХ
# ═══════════════════════════════════════════
add_widget_heading('14. Схема данных (основные типы)', '')

add_section('SalesData (основной тип)')
add_code(
    'interface SalesData {\n'
    '  salesDataByShopName: Record<string, ShopSalesData>\n'
    '  grandTotalSell: number\n'
    '  grandTotalRefund: number\n'
    '  netRevenue: number\n'
    '  totalChecks: number\n'
    '  averageCheck: number\n'
    '  topProducts: ProductData[]\n'
    '  // ...\n'
    '}'
)

add_section('ShopSalesData')
add_code(
    'interface ShopSalesData {\n'
    '  totalSell: number\n'
    '  totalRefund: number\n'
    '  netSell: number\n'
    '  checksCount: number\n'
    '  averageCheck: number\n'
    '  sell: Record<string, number>  // по категориям\n'
    '  refund: Record<string, number>\n'
    '}'
)

add_section('ProductData')
add_code(
    'interface ProductData {\n'
    '  productName: string\n'
    '  netRevenue: number\n'
    '  netQuantity: number\n'
    '  marginPct: number\n'
    '  refundRevenue: number\n'
    '  refundRate: number\n'
    '}'
)

# ═══════════════════════════════════════════
# 15. КОМПОНОВКА СТРАНИЦЫ
# ═══════════════════════════════════════════
add_widget_heading('15. Компоновка домашней страницы', 'Порядок и условия отображения виджетов')

doc.add_paragraph('Страница использует flex-контейнер с центрированием и максимальной шириной:')
add_code('flex flex-col items-center w-full min-h-screen bg-gray-100 dark:bg-gray-900 pt-20 sm:pt-24 px-4 sm:px-6 pb-24')
doc.add_paragraph('Внутренний контейнер: w-full max-w-7xl space-y-4')

add_section('Порядок виджетов (сверху вниз)')
doc.add_paragraph('1. HomeTopBar — верхняя панель (Evo App, индикатор онлайн, кнопка обновления)')
doc.add_paragraph('2. DailyBriefing — всегда')
doc.add_paragraph('3. SellerPerformanceWidget — только SUPERADMIN')
doc.add_paragraph('4. DateFilter — всегда')
doc.add_paragraph('5. PlanStatusWidget — всегда (ИСКЛЮЧЁН из этой спецификации)')
doc.add_paragraph('6. Сетка 2×3 (grid grid-cols-2 gap-4):')
doc.add_paragraph('   a. RevenueWidget — всегда')
doc.add_paragraph('   b. SalesTempoWidget — SUPERADMIN/ADMIN')
doc.add_paragraph('   c. FinanceWidget — SUPERADMIN/ADMIN')
doc.add_paragraph('   d. BestShopWidget — SUPERADMIN/ADMIN')
doc.add_paragraph('   e. TopProductWidget — всегда')
doc.add_paragraph('   f. AccessoriesWidget — всегда (ИСКЛЮЧЁН из этой спецификации)')
doc.add_paragraph('7. TodayAlertsWidget — SUPERADMIN')
doc.add_paragraph('8. StockHealthWidget — SUPERADMIN/ADMIN')
doc.add_paragraph('9. QuickActionsWidget — всегда')
doc.add_paragraph('10. LastUpdated — время последнего обновления')

add_section('Роли и видимость')
add_state_table([
    ('SUPERADMIN', 'Все виджеты', 'Полный доступ'),
    ('ADMIN', 'Revenue, Tempo, Finance, BestShop, TopProduct, StockHealth, QuickActions', 'Без SellerPerformance и Alerts'),
    ('CASHIER', 'Revenue, TopProduct, QuickActions', 'Базовый набор'),
])

# ═══════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════
output_path = '/home/admingimolost/evo-app/docs/home-widgets-spec.docx'
doc.save(output_path)
print(f'Document saved: {output_path}')
